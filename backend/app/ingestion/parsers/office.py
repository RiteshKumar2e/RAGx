"""DOCX, plain-text/Markdown, CSV and HTML parsers."""

from __future__ import annotations

import asyncio
import io
import re
from typing import Any

from app.core.errors import IngestionError
from app.core.logging import get_logger
from app.core.text import normalize_whitespace
from app.ingestion.parsers.base import ContentBlock, DocumentParser, ParsedDocument

log = get_logger("ragx.parse.office")

MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
NUMBERED_HEADING = re.compile(r"^\d+(\.\d+)*\.?\s+[A-Z].{2,90}$")


class DocxParser(DocumentParser):
    """python-docx extraction preserving heading hierarchy and tables."""

    name = "docx"
    extensions = (".docx",)

    async def parse(self, data: bytes, filename: str) -> ParsedDocument:
        return await asyncio.to_thread(self._parse_sync, data, filename)

    def _parse_sync(self, data: bytes, filename: str) -> ParsedDocument:
        try:
            import docx  # noqa: PLC0415
        except ImportError as exc:  # pragma: no cover
            raise IngestionError("python-docx is required to parse .docx files.") from exc

        try:
            document = docx.Document(io.BytesIO(data))
        except Exception as exc:
            raise IngestionError("This DOCX file could not be opened.", detail=str(exc)[:200]) from exc

        parsed = ParsedDocument()
        core = document.core_properties
        parsed.metadata = {
            "author": core.author,
            "created": str(core.created) if core.created else None,
            "last_modified_by": core.last_modified_by,
            "revision": core.revision,
        }
        parsed.title = (core.title or "").strip() or None

        section_stack: list[tuple[int, str]] = []
        order = 0
        buffer: list[str] = []

        def flush() -> None:
            nonlocal order, buffer
            if not buffer:
                return
            text = normalize_whitespace("\n".join(buffer))
            buffer = []
            if not text:
                return
            order += 1
            path = [t for _, t in section_stack]
            parsed.blocks.append(
                ContentBlock(
                    text=text,
                    modality="text",
                    section=path[-1] if path else None,
                    section_path=list(path),
                    order=order,
                )
            )

        for paragraph in document.paragraphs:
            text = normalize_whitespace(paragraph.text)
            if not text:
                continue
            style = (paragraph.style.name or "").lower() if paragraph.style else ""
            level = 0
            if style.startswith("heading"):
                match = re.search(r"(\d+)", style)
                level = int(match.group(1)) if match else 1
            elif style in ("title",):
                level = 1
                parsed.title = parsed.title or text

            if level:
                flush()
                while section_stack and section_stack[-1][0] >= level:
                    section_stack.pop()
                section_stack.append((level, text[:160]))
                parsed.outline.append({"level": level, "title": text, "page": None})
                order += 1
                path = [t for _, t in section_stack]
                parsed.blocks.append(
                    ContentBlock(
                        text=text,
                        modality="text",
                        section=path[-1],
                        section_path=list(path),
                        heading_level=level,
                        order=order,
                    )
                )
            else:
                buffer.append(text)

        flush()

        for table_index, table in enumerate(document.tables, start=1):
            markdown = self._table_to_markdown(table)
            if not markdown:
                continue
            order += 1
            path = [t for _, t in section_stack]
            parsed.blocks.append(
                ContentBlock(
                    text=markdown,
                    modality="table",
                    section=path[-1] if path else None,
                    section_path=list(path),
                    table_label=f"Table {table_index}",
                    order=order,
                    metadata={"rows": len(table.rows), "columns": len(table.columns)},
                )
            )

        parsed.page_count = max(1, round(parsed.text_length / 3000))
        return parsed

    @staticmethod
    def _table_to_markdown(table: Any) -> str:
        rows = [[normalize_whitespace(cell.text) for cell in row.cells] for row in table.rows]
        rows = [r for r in rows if any(r)]
        if len(rows) < 2:
            return ""
        width = max(len(r) for r in rows)
        rows = [r + [""] * (width - len(r)) for r in rows]
        header, *body = rows
        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * width) + " |",
        ]
        lines.extend("| " + " | ".join(r) + " |" for r in body[:60])
        return "\n".join(lines)


class TextParser(DocumentParser):
    """Plain text and Markdown. Markdown headings become the section path."""

    name = "text"
    extensions = (".txt", ".md")

    async def parse(self, data: bytes, filename: str) -> ParsedDocument:
        return await asyncio.to_thread(self._parse_sync, data, filename)

    def _parse_sync(self, data: bytes, filename: str) -> ParsedDocument:
        text = self._decode(data)
        parsed = ParsedDocument(page_count=max(1, round(len(text) / 3000)))
        parsed.metadata = {"characters": len(text), "lines": text.count("\n") + 1}

        section_stack: list[tuple[int, str]] = []
        order = 0
        buffer: list[str] = []
        in_code_fence = False

        def flush(modality: str = "text") -> None:
            nonlocal order, buffer
            if not buffer:
                return
            body = "\n".join(buffer).strip()
            buffer = []
            if not body:
                return
            order += 1
            path = [t for _, t in section_stack]
            parsed.blocks.append(
                ContentBlock(
                    text=body if modality == "code" else normalize_whitespace(body),
                    modality=modality,
                    section=path[-1] if path else None,
                    section_path=list(path),
                    order=order,
                )
            )

        for line in text.splitlines():
            if line.strip().startswith("```"):
                if in_code_fence:
                    buffer.append(line)
                    flush("code")
                    in_code_fence = False
                else:
                    flush()
                    buffer.append(line)
                    in_code_fence = True
                continue
            if in_code_fence:
                buffer.append(line)
                continue

            heading = MD_HEADING.match(line)
            if heading or (line.strip() and NUMBERED_HEADING.match(line.strip())):
                flush()
                if heading:
                    level, title = len(heading.group(1)), heading.group(2).strip()
                else:
                    level, title = 2, line.strip()
                while section_stack and section_stack[-1][0] >= level:
                    section_stack.pop()
                section_stack.append((level, title[:160]))
                parsed.outline.append({"level": level, "title": title, "page": None})
                order += 1
                path = [t for _, t in section_stack]
                parsed.blocks.append(
                    ContentBlock(
                        text=title,
                        modality="text",
                        section=path[-1],
                        section_path=list(path),
                        heading_level=level,
                        order=order,
                    )
                )
                continue

            if not line.strip():
                flush()
            else:
                buffer.append(line)

        flush("code" if in_code_fence else "text")
        if not parsed.title and parsed.outline:
            parsed.title = parsed.outline[0]["title"]
        return parsed

    @staticmethod
    def _decode(data: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")


class CSVParser(DocumentParser):
    """CSV parser.

    A CSV is inherently tabular, so it is emitted as table blocks: a schema
    summary (column names, dtypes, describe()) plus row windows rendered as
    Markdown, which keeps numeric questions answerable and citable by row range.
    """

    name = "csv"
    extensions = (".csv",)
    ROWS_PER_BLOCK = 25
    MAX_BLOCKS = 40

    async def parse(self, data: bytes, filename: str) -> ParsedDocument:
        return await asyncio.to_thread(self._parse_sync, data, filename)

    def _parse_sync(self, data: bytes, filename: str) -> ParsedDocument:
        try:
            import pandas as pd  # noqa: PLC0415
        except ImportError as exc:  # pragma: no cover
            raise IngestionError("pandas is required to parse CSV files.") from exc

        try:
            frame = pd.read_csv(io.BytesIO(data))
        except Exception:
            try:
                frame = pd.read_csv(io.BytesIO(data), sep=None, engine="python", on_bad_lines="skip")
            except Exception as exc:
                raise IngestionError("This CSV could not be parsed.", detail=str(exc)[:200]) from exc

        parsed = ParsedDocument(page_count=1)
        parsed.title = filename
        parsed.metadata = {
            "rows": int(frame.shape[0]),
            "columns": int(frame.shape[1]),
            "column_names": [str(c) for c in frame.columns],
            "dtypes": {str(c): str(t) for c, t in frame.dtypes.items()},
        }

        summary_lines = [
            f"Dataset: {filename}",
            f"Rows: {frame.shape[0]}, Columns: {frame.shape[1]}",
            "Columns: " + ", ".join(f"{c} ({frame.dtypes[c]})" for c in frame.columns),
        ]
        numeric = frame.select_dtypes(include="number")
        if not numeric.empty:
            try:
                summary_lines.append("\nNumeric summary:\n" + numeric.describe().to_markdown())
            except Exception:
                summary_lines.append("\nNumeric summary:\n" + numeric.describe().to_string())

        parsed.blocks.append(
            ContentBlock(
                text="\n".join(summary_lines),
                modality="table",
                section="Schema",
                section_path=["Schema"],
                table_label="Schema",
                order=0,
                metadata={"kind": "schema"},
            )
        )

        for order, start in enumerate(range(0, len(frame), self.ROWS_PER_BLOCK), start=1):
            if order > self.MAX_BLOCKS:
                parsed.warnings.append(
                    f"Only the first {self.MAX_BLOCKS * self.ROWS_PER_BLOCK} rows were indexed."
                )
                break
            window = frame.iloc[start : start + self.ROWS_PER_BLOCK]
            try:
                markdown = window.to_markdown(index=False)
            except Exception:
                markdown = window.to_string(index=False)
            end = start + len(window)
            parsed.blocks.append(
                ContentBlock(
                    text=f"Rows {start + 1}-{end} of {filename}:\n\n{markdown}",
                    modality="table",
                    section=f"Rows {start + 1}-{end}",
                    section_path=["Data"],
                    table_label=f"Rows {start + 1}-{end}",
                    order=order,
                    metadata={"row_start": start + 1, "row_end": end, "kind": "rows"},
                )
            )
        return parsed


class HTMLParser(DocumentParser):
    """HTML / fetched web content."""

    name = "html"
    extensions = (".html", ".htm")

    async def parse(self, data: bytes, filename: str) -> ParsedDocument:
        return await asyncio.to_thread(self._parse_sync, data, filename)

    def _parse_sync(self, data: bytes, filename: str) -> ParsedDocument:
        from bs4 import BeautifulSoup  # noqa: PLC0415

        soup = BeautifulSoup(data, "lxml")
        for tag in soup(["script", "style", "nav", "footer", "aside", "noscript"]):
            tag.decompose()

        parsed = ParsedDocument(page_count=1)
        parsed.title = (soup.title.string or "").strip() if soup.title and soup.title.string else None
        description = soup.find("meta", attrs={"name": "description"})
        parsed.metadata = {
            "description": description.get("content") if description else None,
            "source": filename,
        }

        section_stack: list[tuple[int, str]] = []
        order = 0
        main = soup.find("article") or soup.find("main") or soup.body or soup

        for element in main.find_all(["h1", "h2", "h3", "h4", "p", "li", "pre", "table"]):
            text = normalize_whitespace(element.get_text(" ", strip=True))
            if not text or len(text) < 3:
                continue
            name = element.name
            if name.startswith("h"):
                level = int(name[1])
                while section_stack and section_stack[-1][0] >= level:
                    section_stack.pop()
                section_stack.append((level, text[:160]))
                parsed.outline.append({"level": level, "title": text, "page": None})
                continue

            order += 1
            path = [t for _, t in section_stack]
            modality = "table" if name == "table" else ("code" if name == "pre" else "text")
            parsed.blocks.append(
                ContentBlock(
                    text=text,
                    modality=modality,
                    section=path[-1] if path else None,
                    section_path=list(path),
                    order=order,
                    table_label=f"Table {sum(1 for b in parsed.blocks if b.modality == 'table') + 1}"
                    if modality == "table"
                    else None,
                )
            )
        return parsed


__all__ = ["DocxParser", "TextParser", "CSVParser", "HTMLParser"]
